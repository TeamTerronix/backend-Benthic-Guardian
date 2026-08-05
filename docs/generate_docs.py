"""
generate_docs.py
================
Generates all diagram PNGs and then builds the Word documentation file.
Run from the backend/ directory:
    python docs/generate_docs.py
"""

import os
import sys
import textwrap
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.patches as patches
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
import numpy as np

# ── Output dirs ──────────────────────────────────────────────────────────────
DOCS_DIR = Path(__file__).parent
IMG_DIR  = DOCS_DIR / "images"
IMG_DIR.mkdir(parents=True, exist_ok=True)

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def save(fig, name):
    path = IMG_DIR / name
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    print(f"  saved {path}")
    return path


def box(ax, x, y, w, h, text, bg="#2563EB", fg="white", fontsize=9,
        radius=0.02, bold=False, wrap=False):
    fancy = FancyBboxPatch((x - w/2, y - h/2), w, h,
                            boxstyle=f"round,pad=0.01,rounding_size={radius}",
                            linewidth=1.2, edgecolor="#1e3a5f",
                            facecolor=bg, zorder=3)
    ax.add_patch(fancy)
    weight = "bold" if bold else "normal"
    lines = textwrap.wrap(text, 18) if wrap else [text]
    ax.text(x, y, "\n".join(lines), ha="center", va="center",
            fontsize=fontsize, color=fg, fontweight=weight, zorder=4,
            multialignment="center")


def arrow(ax, x1, y1, x2, y2, label="", color="#374151", lw=1.5, style="->"):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle=style, color=color, lw=lw,
                                connectionstyle="arc3,rad=0.0"),
                zorder=2)
    if label:
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx+0.02, my+0.02, label, fontsize=7, color="#374151", zorder=5)


def dashed_arrow(ax, x1, y1, x2, y2, label="", color="#6B7280"):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle="->", color=color, lw=1.2,
                                linestyle="dashed",
                                connectionstyle="arc3,rad=0.0"),
                zorder=2)
    if label:
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx+0.02, my+0.02, label, fontsize=7, color=color, zorder=5,
                style="italic")


# ═════════════════════════════════════════════════════════════════════════════
# DIAGRAM 1 — System Context (C4 L1)
# ═════════════════════════════════════════════════════════════════════════════

def diagram_system_context():
    fig, ax = plt.subplots(figsize=(13, 7))
    ax.set_xlim(0, 13); ax.set_ylim(0, 7)
    ax.axis("off")
    fig.patch.set_facecolor("#F8FAFC")
    ax.set_facecolor("#F8FAFC")

    ax.text(6.5, 6.65, "SLIOT — System Context", fontsize=14, fontweight="bold",
            ha="center", va="center", color="#1e3a5f")
    ax.text(6.5, 6.35, "How users, devices, and services connect to the SLIOT Backend",
            fontsize=9, ha="center", va="center", color="#6B7280")

    # Users
    box(ax, 1.4, 5.2, 2.0, 0.7, "Researcher / Admin\n(Dashboard user)", bg="#6366F1")
    box(ax, 1.4, 3.7, 2.0, 0.7, "ESP32 Sensor\n(IoT device)", bg="#059669")

    # Central system
    box(ax, 6.5, 4.5, 3.2, 1.8,
        "SLIOT Backend\n(FastAPI · Docker · Render)",
        bg="#1D4ED8", fontsize=11, bold=True)

    # External systems
    box(ax, 11.0, 5.5, 2.2, 0.7, "Supabase\n(PostgreSQL)", bg="#0F766E")
    box(ax, 11.0, 4.5, 2.2, 0.7, "Vercel\n(Next.js Dashboard)", bg="#7C3AED")
    box(ax, 11.0, 3.5, 2.2, 0.7, "GitHub Actions\n(CI / CD)", bg="#374151")

    # Arrows — users to backend
    arrow(ax, 2.4, 5.2,  5.0, 4.9, "HTTPS + JWT")
    arrow(ax, 2.4, 3.7,  5.0, 4.2, "POST /data")

    # Backend to external
    arrow(ax, 8.1, 4.8,  9.9, 5.5,  "Reads & writes")
    arrow(ax, 8.1, 4.5,  9.9, 4.5,  "API responses")
    arrow(ax, 8.1, 4.2,  9.9, 3.5,  "CI on push/PR")

    # Legend
    legend_items = [
        mpatches.Patch(color="#6366F1", label="External Users"),
        mpatches.Patch(color="#059669", label="IoT Devices"),
        mpatches.Patch(color="#1D4ED8", label="This System (Backend)"),
        mpatches.Patch(color="#0F766E", label="External Services"),
    ]
    ax.legend(handles=legend_items, loc="lower left", fontsize=8,
              framealpha=0.85, edgecolor="#CBD5E1")

    return save(fig, "01_system_context.png")


# ═════════════════════════════════════════════════════════════════════════════
# DIAGRAM 2 — Container / Component Architecture
# ═════════════════════════════════════════════════════════════════════════════

def diagram_container():
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.set_xlim(0, 14); ax.set_ylim(0, 8)
    ax.axis("off")
    fig.patch.set_facecolor("#F8FAFC")
    ax.set_facecolor("#F8FAFC")

    ax.text(7, 7.65, "SLIOT — Backend Container Architecture", fontsize=13,
            fontweight="bold", ha="center", va="center", color="#1e3a5f")

    # ── Left: external actors ──
    box(ax, 1.2, 6.0, 1.8, 0.6, "Dashboard\n(Vercel)", bg="#7C3AED", fontsize=8)
    box(ax, 1.2, 5.0, 1.8, 0.6, "ESP32 Sensor", bg="#059669", fontsize=8)
    box(ax, 1.2, 4.0, 1.8, 0.6, "Admin User", bg="#6366F1", fontsize=8)

    # ── FastAPI App (big box) ──
    system_rect = FancyBboxPatch((2.8, 2.0), 7.2, 5.0,
                                  boxstyle="round,pad=0.05",
                                  linewidth=2, edgecolor="#1D4ED8",
                                  facecolor="#EFF6FF", zorder=1)
    ax.add_patch(system_rect)
    ax.text(6.4, 6.8, "FastAPI Application  (main.py)", fontsize=9,
            fontweight="bold", ha="center", color="#1D4ED8", zorder=3)

    # Components inside
    box(ax, 4.5, 6.1, 2.0, 0.65, "Auth Module\n(auth.py · JWT · bcrypt)", bg="#2563EB", fontsize=7.5)
    box(ax, 7.0, 6.1, 2.0, 0.65, "CORS Middleware", bg="#3B82F6", fontsize=7.5)
    box(ax, 4.5, 5.1, 2.0, 0.65, "REST API Routes\n(/data · /sensors · /auth)", bg="#1D4ED8", fontsize=7.5)
    box(ax, 7.0, 5.1, 2.0, 0.65, "WebSocket Manager\n(/ws/alerts)", bg="#1D4ED8", fontsize=7.5)
    box(ax, 4.5, 4.1, 2.0, 0.65, "APScheduler\n(6h PINN · 24h archive)", bg="#6366F1", fontsize=7.5)
    box(ax, 7.0, 4.1, 2.0, 0.65, "DB Layer\n(database.py · SQLAlchemy)", bg="#0891B2", fontsize=7.5)
    box(ax, 4.5, 3.1, 2.0, 0.65, "PINN Forecaster\n(forecaster.py · TF)", bg="#7C3AED", fontsize=7.5)
    box(ax, 7.0, 3.1, 2.0, 0.65, "ANN-LSTM\n(lstm_forecaster.py · TF)", bg="#7C3AED", fontsize=7.5)

    # ── Right: external services ──
    box(ax, 12.0, 6.0, 2.0, 0.65, "Supabase\n(PostgreSQL)", bg="#0F766E", fontsize=8)
    box(ax, 12.0, 4.5, 2.0, 0.65, "model/\nPINN + LSTM .h5", bg="#92400E", fontsize=8)
    box(ax, 12.0, 3.0, 2.0, 0.65, "history_archive.csv", bg="#374151", fontsize=8)

    # Arrows — actors to API
    arrow(ax, 2.1, 6.0, 2.8, 5.8,  "HTTP")
    arrow(ax, 2.1, 5.0, 2.8, 5.0,  "POST /data")
    arrow(ax, 2.1, 4.0, 2.8, 4.3,  "HTTP")

    # Internal wiring
    arrow(ax, 5.5, 5.4, 5.5, 4.8)
    arrow(ax, 8.0, 5.4, 8.0, 4.8)
    arrow(ax, 5.5, 4.4, 5.5, 3.8)
    arrow(ax, 8.0, 4.4, 8.0, 3.8)

    # DB to Supabase
    arrow(ax, 9.0, 4.1, 11.0, 5.7,  "SQL/SSL")
    # Models from disk
    arrow(ax, 5.5, 2.8, 11.0, 4.5,  "load .h5")
    arrow(ax, 7.0, 2.8, 11.0, 4.5)
    # Archive
    arrow(ax, 5.5, 3.8, 11.0, 3.0,  "append CSV")

    return save(fig, "02_container_architecture.png")


# ═════════════════════════════════════════════════════════════════════════════
# DIAGRAM 3 — Entity Relationship Diagram (ERD)
# ═════════════════════════════════════════════════════════════════════════════

def diagram_erd():
    fig, ax = plt.subplots(figsize=(15, 9))
    ax.set_xlim(0, 15); ax.set_ylim(0, 9)
    ax.axis("off")
    fig.patch.set_facecolor("#FAFAF9")
    ax.set_facecolor("#FAFAF9")

    ax.text(7.5, 8.65, "SLIOT — Database Entity Relationship Diagram",
            fontsize=13, fontweight="bold", ha="center", color="#1e3a5f")

    def table(ax, x, y, title, fields, w=2.6, row_h=0.28, header_bg="#1D4ED8"):
        header_h = 0.38
        total_h  = header_h + len(fields) * row_h

        # Header
        hdr = FancyBboxPatch((x, y - header_h), w, header_h,
                              boxstyle="round,pad=0.01,rounding_size=0.02",
                              linewidth=1.5, edgecolor="#1e3a5f",
                              facecolor=header_bg, zorder=3)
        ax.add_patch(hdr)
        ax.text(x + w/2, y - header_h/2, title, ha="center", va="center",
                fontsize=9, fontweight="bold", color="white", zorder=4)

        # Rows
        for i, (col, typ, note) in enumerate(fields):
            ry = y - header_h - (i+1)*row_h
            bg = "#DBEAFE" if i % 2 == 0 else "#EFF6FF"
            row_patch = patches.Rectangle((x, ry), w, row_h,
                                           linewidth=0.5, edgecolor="#CBD5E1",
                                           facecolor=bg, zorder=2)
            ax.add_patch(row_patch)
            prefix = "[PK] " if "PK" in note else ("[FK] " if "FK" in note else "     ")
            ax.text(x + 0.07, ry + row_h/2,
                    f"{prefix}{col}",
                    ha="left", va="center", fontsize=7.2, color="#111827", zorder=4)
            ax.text(x + w - 0.07, ry + row_h/2,
                    typ,
                    ha="right", va="center", fontsize=6.5, color="#6B7280", zorder=4)

        # Outer border
        border = patches.Rectangle((x, y - header_h - len(fields)*row_h),
                                     w, total_h,
                                     linewidth=1.5, edgecolor="#1e3a5f",
                                     facecolor="none", zorder=5)
        ax.add_patch(border)
        return (x + w/2, y - header_h - len(fields)*row_h)  # bottom center

    # ── users ──
    table(ax, 0.4, 8.2, "users", [
        ("id",              "INTEGER",  "PK"),
        ("email",           "VARCHAR",  "UNIQUE"),
        ("hashed_password", "VARCHAR",  ""),
        ("role",            "ENUM",     "admin/user"),
        ("created_at",      "TIMESTAMPTZ", ""),
    ], header_bg="#1D4ED8")

    # ── network_groups ──
    table(ax, 4.0, 8.2, "network_groups", [
        ("id",         "VARCHAR",     "PK  e.g. ng_xxx"),
        ("name",       "VARCHAR",     "nullable"),
        ("created_at", "TIMESTAMPTZ", ""),
    ], header_bg="#0891B2")

    # ── user_network_groups ──
    table(ax, 7.8, 8.2, "user_network_groups", [
        ("user_id",          "INTEGER", "PK · FK → users"),
        ("network_group_id", "VARCHAR", "PK · FK → network_groups"),
        ("created_at",       "TIMESTAMPTZ", ""),
    ], header_bg="#6366F1")

    # ── sensors ──
    table(ax, 0.4, 5.4, "sensors", [
        ("id",               "INTEGER", "PK"),
        ("sensor_uid",       "VARCHAR", "UNIQUE  device ID"),
        ("owner_id",         "INTEGER", "FK → users"),
        ("network_group_id", "VARCHAR", "FK → network_groups"),
        ("latitude",         "FLOAT",   ""),
        ("longitude",        "FLOAT",   ""),
        ("depth",            "FLOAT",   ""),
        ("is_approved",      "BOOLEAN", "default false"),
        ("created_at",       "TIMESTAMPTZ", ""),
    ], header_bg="#059669")

    # ── sensor_readings ──
    table(ax, 4.0, 5.4, "sensor_readings", [
        ("id",          "INTEGER",     "PK"),
        ("sensor_id",   "INTEGER",     "FK → sensors"),
        ("timestamp",   "TIMESTAMPTZ", ""),
        ("temperature", "FLOAT",       "°C"),
        ("created_at",  "TIMESTAMPTZ", ""),
    ], header_bg="#D97706")

    # ── predictions ──
    table(ax, 8.8, 5.4, "predictions", [
        ("id",               "INTEGER",     "PK"),
        ("sensor_id",        "INTEGER",     "FK → sensors"),
        ("target_timestamp", "TIMESTAMPTZ", ""),
        ("predicted_temp",   "FLOAT",       "°C"),
        ("risk_level",       "INTEGER",     "0=ok 1=warn 2=danger"),
        ("risk_score",       "FLOAT",       "0–1 continuous"),
        ("anomaly",          "FLOAT",       "°C above baseline"),
        ("days_stressed",    "INTEGER",     "consecutive warm days"),
        ("warming_rate",     "FLOAT",       "°C/day"),
        ("physics_residual", "FLOAT",       "|R|² heat eq."),
        ("created_at",       "TIMESTAMPTZ", ""),
    ], header_bg="#7C3AED")

    # Relationship lines
    def rel_line(ax, x1, y1, x2, y2, label=""):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color="#374151", lw=1.3),
                    zorder=6)
        if label:
            mx, my = (x1+x2)/2, (y1+y2)/2
            ax.text(mx+0.05, my, label, fontsize=6.5, color="#374151", zorder=7)

    # users → user_network_groups
    rel_line(ax, 1.7,  6.65, 7.8,  7.4,  "1:N")
    # network_groups → user_network_groups
    rel_line(ax, 6.66, 7.0,  7.8,  7.1,  "1:N")
    # users → sensors
    rel_line(ax, 1.7,  5.4,  1.7,  4.55, "1:N")
    # network_groups → sensors
    rel_line(ax, 5.3,  6.55, 2.5,  4.3,  "1:N")
    # sensors → sensor_readings
    rel_line(ax, 3.0,  3.8,  4.0,  4.2,  "1:N")
    # sensors → predictions
    rel_line(ax, 3.0,  3.5,  8.8,  4.0,  "1:N")

    # Legend
    legend_items = [
        mpatches.Patch(color="#1D4ED8", label="users"),
        mpatches.Patch(color="#0891B2", label="network_groups"),
        mpatches.Patch(color="#6366F1", label="user_network_groups (join)"),
        mpatches.Patch(color="#059669", label="sensors"),
        mpatches.Patch(color="#D97706", label="sensor_readings"),
        mpatches.Patch(color="#7C3AED", label="predictions"),
    ]
    ax.legend(handles=legend_items, loc="lower left", fontsize=7.5,
              framealpha=0.9, edgecolor="#CBD5E1", ncol=3)

    return save(fig, "03_erd.png")


# ═════════════════════════════════════════════════════════════════════════════
# DIAGRAM 4 — Request/Response Sequence (login → ingest → forecast)
# ═════════════════════════════════════════════════════════════════════════════

def diagram_sequence():
    fig, ax = plt.subplots(figsize=(14, 9))
    ax.set_xlim(0, 14); ax.set_ylim(0, 9)
    ax.axis("off")
    fig.patch.set_facecolor("#FAFAF9")
    ax.set_facecolor("#FAFAF9")

    ax.text(7, 8.7, "SLIOT — Key Request Sequences", fontsize=13,
            fontweight="bold", ha="center", color="#1e3a5f")

    actors = [
        (1.2,  "ESP32\nSensor",   "#059669"),
        (3.5,  "Dashboard\n(User)", "#7C3AED"),
        (6.5,  "FastAPI\nBackend", "#1D4ED8"),
        (9.5,  "Database\n(Supabase)", "#0F766E"),
        (12.5, "PINN/LSTM\nModels",    "#92400E"),
    ]

    # Draw actor boxes + lifelines
    for x, label, col in actors:
        box(ax, x, 8.2, 1.6, 0.55, label, bg=col, fontsize=8)
        ax.plot([x, x], [7.93, 0.2], color="#9CA3AF", lw=1, ls="--", zorder=1)

    def seq_arrow(y, x1, x2, label, response=False, color="#374151"):
        ax.annotate("", xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle="->", color=color, lw=1.3,
                                    linestyle="dashed" if response else "solid"),
                    zorder=3)
        mid = (x1 + x2) / 2
        ax.text(mid, y + 0.1, label, ha="center", va="bottom",
                fontsize=7.5, color=color, zorder=4,
                style="italic" if response else "normal")

    def section(y, text):
        ax.add_patch(patches.FancyBboxPatch((0.1, y - 0.12), 13.8, 0.25,
                      boxstyle="round,pad=0.02", facecolor="#E0E7FF",
                      edgecolor="#6366F1", lw=0.8, zorder=0))
        ax.text(0.3, y, text, fontsize=8, fontweight="bold", color="#3730A3", va="center")

    # ── Section A: Auth ──────────────────────────────────────────
    section(7.65, "A  User Login")
    seq_arrow(7.3,  3.5,  6.5, "POST /auth/token  {email, password}")
    seq_arrow(7.0,  6.5,  9.5, "SELECT user WHERE email = ?")
    seq_arrow(6.75, 9.5,  6.5, "user row",  response=True)
    seq_arrow(6.5,  6.5,  3.5, "200 { access_token: JWT }",  response=True)

    # ── Section B: Sensor ingest ─────────────────────────────────
    section(6.15, "B  Sensor Data Ingest  (POST /data)")
    seq_arrow(5.85, 1.2,  6.5, "POST /data  {sensor_uid, temperature}")
    seq_arrow(5.6,  6.5,  9.5, "SELECT sensor WHERE uid = ?")
    seq_arrow(5.35, 9.5,  6.5, "sensor row", response=True)
    seq_arrow(5.1,  6.5,  9.5, "INSERT sensor_readings")
    seq_arrow(4.85, 9.5,  6.5, "ok",          response=True)
    seq_arrow(4.6,  6.5,  1.2, "200 { status: created }", response=True)
    seq_arrow(4.35, 6.5, 12.5, "spawn forecast thread (throttled 45s)")

    # ── Section C: Dashboard reads PINN forecast ─────────────────
    section(4.0, "C  Dashboard Reads PINN Forecast")
    seq_arrow(3.7,  3.5,  6.5, "GET /sensors/{id}/forecast  (JWT)")
    seq_arrow(3.45, 6.5,  9.5, "SELECT predictions WHERE sensor_id = ?")
    seq_arrow(3.2,  9.5,  6.5, "168 rows (7 days × 24 h)", response=True)
    seq_arrow(2.95, 6.5,  3.5, "[ {timestamp, predicted_temp, risk_level} ]", response=True)

    # ── Section D: LSTM Forecast ──────────────────────────────────
    section(2.6, "D  ANN-LSTM Forecast")
    seq_arrow(2.3,  3.5,  6.5, "GET /api/lstm-forecast?location=hikkaduwa  (JWT)")
    seq_arrow(2.05, 6.5, 12.5, "load model, run 60-day history → +1/+3/+7d")
    seq_arrow(1.8, 12.5,  6.5, "forecast points", response=True)
    seq_arrow(1.55, 6.5,  3.5, "[ { horizon_days, predicted_temp, risk } ]", response=True)

    return save(fig, "04_sequence.png")


# ═════════════════════════════════════════════════════════════════════════════
# DIAGRAM 5 — Scheduler Timeline
# ═════════════════════════════════════════════════════════════════════════════

def diagram_scheduler():
    fig, ax = plt.subplots(figsize=(13, 5))
    ax.set_xlim(-0.5, 25); ax.set_ylim(0, 5)
    ax.axis("off")
    fig.patch.set_facecolor("#F8FAFC")
    ax.set_facecolor("#F8FAFC")

    ax.text(12, 4.7, "SLIOT — Background Scheduler Jobs (24-hour view)",
            fontsize=12, fontweight="bold", ha="center", color="#1e3a5f")

    # Timeline axis
    ax.plot([-0.2, 24.5], [2.5, 2.5], color="#374151", lw=1.5)
    for h in range(0, 25):
        ax.plot([h, h], [2.4, 2.6], color="#374151", lw=0.8)
        if h % 6 == 0:
            ax.text(h, 2.2, f"{h:02d}:00", ha="center", fontsize=7.5, color="#374151")

    ax.text(12, 1.85, "Time (UTC hours)", ha="center", fontsize=8, color="#6B7280")

    # Job 1: PINN every 6h  (0,6,12,18)
    for h in [0, 6, 12, 18]:
        rect = patches.FancyBboxPatch((h, 3.0), 1.0, 0.65,
                                       boxstyle="round,pad=0.03",
                                       facecolor="#7C3AED", edgecolor="#4C1D95",
                                       lw=1, zorder=3)
        ax.add_patch(rect)
        ax.text(h + 0.5, 3.32, "PINN\nforecast", ha="center", va="center",
                fontsize=6.5, color="white", fontweight="bold")
        ax.plot([h+0.5, h+0.5], [3.0, 2.6], color="#7C3AED", lw=1, ls=":")

    ax.text(-0.4, 3.32, "Job 1\n(6h)", ha="right", va="center",
            fontsize=7.5, color="#7C3AED", fontweight="bold")

    # Job 2: archive every 24h (at 00:00)
    rect2 = patches.FancyBboxPatch((0, 1.5), 1.2, 0.65,
                                    boxstyle="round,pad=0.03",
                                    facecolor="#D97706", edgecolor="#92400E",
                                    lw=1, zorder=3)
    ax.add_patch(rect2)
    ax.text(0.6, 1.82, "Archive\n(>30d)", ha="center", va="center",
            fontsize=6.5, color="white", fontweight="bold")
    ax.plot([0.6, 0.6], [1.5, 2.6], color="#D97706", lw=1, ls=":")

    ax.text(-0.4, 1.82, "Job 2\n(24h)", ha="right", va="center",
            fontsize=7.5, color="#D97706", fontweight="bold")

    # On-demand trigger
    ax.annotate("", xy=(2.5, 2.6), xytext=(2.5, 4.0),
                arrowprops=dict(arrowstyle="->", color="#059669", lw=1.5))
    ax.text(2.5, 4.1, "POST /data triggers\nPINN forecast\n(throttled ≥45s)",
            ha="center", va="bottom", fontsize=7.5, color="#059669",
            bbox=dict(boxstyle="round,pad=0.3", facecolor="#D1FAE5",
                      edgecolor="#059669", lw=1))

    legend_items = [
        mpatches.Patch(color="#7C3AED", label="Job 1 · PINN forecast (every 6h)  →  updates predictions table"),
        mpatches.Patch(color="#D97706", label="Job 2 · Data archive (every 24h)  →  moves >30-day readings to CSV"),
        mpatches.Patch(color="#059669", label="On-demand · POST /data triggers forecast (45s cooldown)"),
    ]
    ax.legend(handles=legend_items, loc="upper right", fontsize=7.5,
              framealpha=0.9, edgecolor="#CBD5E1")

    return save(fig, "05_scheduler.png")


# ═════════════════════════════════════════════════════════════════════════════
# DIAGRAM 6 — Auth / JWT Flow
# ═════════════════════════════════════════════════════════════════════════════

def diagram_auth():
    fig, ax = plt.subplots(figsize=(11, 6))
    ax.set_xlim(0, 11); ax.set_ylim(0, 6)
    ax.axis("off")
    fig.patch.set_facecolor("#F8FAFC")
    ax.set_facecolor("#F8FAFC")

    ax.text(5.5, 5.7, "SLIOT — Authentication & Authorization Flow",
            fontsize=12, fontweight="bold", ha="center", color="#1e3a5f")

    # Boxes
    box(ax, 1.5, 4.5, 2.0, 0.65, "Client\n(Dashboard / ESP32)", bg="#6366F1")
    box(ax, 5.5, 4.5, 2.2, 0.65, "POST /auth/token\n(OAuth2 password flow)", bg="#1D4ED8")
    box(ax, 9.3, 4.5, 1.8, 0.65, "Database\n(users table)", bg="#0F766E")
    box(ax, 5.5, 3.1, 2.2, 0.65, "bcrypt.checkpw()\nverify password", bg="#0891B2")
    box(ax, 5.5, 1.9, 2.2, 0.65, "create_access_token()\nJWT (HS256 · 60 min)", bg="#7C3AED")
    box(ax, 1.5, 2.5, 2.0, 0.65, "Subsequent Requests\nAuthorization: Bearer <token>", bg="#6366F1", fontsize=7.5)
    box(ax, 5.5, 0.7, 2.2, 0.65, "get_current_user()\ndecode JWT → user row", bg="#1D4ED8")

    arrow(ax, 2.5,  4.5, 4.4,  4.5, "{email, password}")
    arrow(ax, 6.6,  4.5, 8.4,  4.5, "SELECT WHERE email")
    arrow(ax, 8.4,  4.3, 6.6,  3.4, "hashed_password")
    arrow(ax, 5.5,  3.8, 5.5,  3.4)
    arrow(ax, 5.5,  2.8, 5.5,  2.2)
    arrow(ax, 4.4,  1.9, 2.5,  2.5, "JWT token", color="#7C3AED")
    arrow(ax, 2.5,  2.2, 4.4,  1.0, "Bearer token")
    arrow(ax, 6.6,  0.7, 8.4,  0.7)
    ax.text(7.5, 0.82, "SELECT user", fontsize=7, color="#374151")

    # Role note
    ax.add_patch(patches.FancyBboxPatch((0.2, 0.2), 3.0, 0.75,
                  boxstyle="round,pad=0.05", facecolor="#FEF3C7",
                  edgecolor="#D97706", lw=1))
    ax.text(1.7, 0.57,
            "Role check:\nadmin → full access\nuser  → own network only",
            ha="center", va="center", fontsize=7, color="#92400E")

    return save(fig, "06_auth_flow.png")


# ═════════════════════════════════════════════════════════════════════════════
# DIAGRAM 7 — ML Models Overview
# ═════════════════════════════════════════════════════════════════════════════

def diagram_ml():
    fig, ax = plt.subplots(figsize=(13, 7))
    ax.set_xlim(0, 13); ax.set_ylim(0, 7)
    ax.axis("off")
    fig.patch.set_facecolor("#F8FAFC")
    ax.set_facecolor("#F8FAFC")

    ax.text(6.5, 6.7, "SLIOT — ML Models in the Backend",
            fontsize=13, fontweight="bold", ha="center", color="#1e3a5f")

    # ── PINN ──
    pinn_rect = FancyBboxPatch((0.3, 1.0), 5.6, 5.0,
                                boxstyle="round,pad=0.05",
                                linewidth=2, edgecolor="#7C3AED",
                                facecolor="#FAF5FF", zorder=1)
    ax.add_patch(pinn_rect)
    ax.text(3.1, 5.7, "PINN  (Physics-Informed Neural Network)",
            fontsize=9.5, fontweight="bold", ha="center", color="#7C3AED", zorder=3)

    box(ax, 3.1, 5.0, 4.5, 0.5, "Inputs: (lat, lon, time) per sensor", bg="#7C3AED", fontsize=8)
    box(ax, 3.1, 4.3, 4.5, 0.5, "Output: temperature every hour → 168 h ahead", bg="#6D28D9", fontsize=8)
    box(ax, 3.1, 3.6, 4.5, 0.5, "Physics: heat-equation PDE residual in loss", bg="#5B21B6", fontsize=8)
    box(ax, 3.1, 2.9, 4.5, 0.5, "Triggered: every 6 h (scheduler) OR after POST /data", bg="#4C1D95", fontsize=8)
    box(ax, 3.1, 2.2, 4.5, 0.5, "Stored in: predictions table (DB)", bg="#3B0764", fontsize=8)
    box(ax, 3.1, 1.5, 4.5, 0.35, "File: model/pinn_model_best.h5 + scalers.pkl", bg="#581C87", fontsize=7.5)

    # ── ANN-LSTM ──
    lstm_rect = FancyBboxPatch((7.0, 1.0), 5.6, 5.0,
                                boxstyle="round,pad=0.05",
                                linewidth=2, edgecolor="#D97706",
                                facecolor="#FFFBEB", zorder=1)
    ax.add_patch(lstm_rect)
    ax.text(9.8, 5.7, "ANN-LSTM  (60-day history forecaster)",
            fontsize=9.5, fontweight="bold", ha="center", color="#D97706", zorder=3)

    box(ax, 9.8, 5.0, 4.5, 0.5, "Inputs: last 60 days SST + DHW per reef site", bg="#D97706", fontsize=8)
    box(ax, 9.8, 4.3, 4.5, 0.5, "Output: SST at +1, +3, +7 days ahead", bg="#B45309", fontsize=8)
    box(ax, 9.8, 3.6, 4.5, 0.5, "Lookback: 60-day sliding window", bg="#92400E", fontsize=8)
    box(ax, 9.8, 2.9, 4.5, 0.5, "Triggered: on API request (lazy-load)", bg="#78350F", fontsize=8)
    box(ax, 9.8, 2.2, 4.5, 0.5, "Source: bundled CSV history (not live DB)", bg="#451A03", fontsize=8)
    box(ax, 9.8, 1.5, 4.5, 0.35, "File: model/ann_lstm_L60_best.h5", bg="#713F12", fontsize=7.5)

    # Comparison axis
    ax.add_patch(patches.FancyBboxPatch((0.3, 0.2), 12.4, 0.55,
                  boxstyle="round,pad=0.04", facecolor="#F0FDF4",
                  edgecolor="#059669", lw=1.2))
    ax.text(6.5, 0.47,
            "Both use TensorFlow · Complementary: PINN → spatial maps & hourly sensor forecast  |  LSTM → short-term SST/DHW sequence",
            ha="center", va="center", fontsize=7.5, color="#059669")

    return save(fig, "07_ml_models.png")


# ═════════════════════════════════════════════════════════════════════════════
# DIAGRAM 8 — Deployment Architecture
# ═════════════════════════════════════════════════════════════════════════════

def diagram_deployment():
    fig, ax = plt.subplots(figsize=(13, 6.5))
    ax.set_xlim(0, 13); ax.set_ylim(0, 6.5)
    ax.axis("off")
    fig.patch.set_facecolor("#F8FAFC")
    ax.set_facecolor("#F8FAFC")

    ax.text(6.5, 6.2, "SLIOT — Production Deployment Architecture",
            fontsize=13, fontweight="bold", ha="center", color="#1e3a5f")

    # IoT
    box(ax, 1.2, 4.8, 1.8, 0.65, "ESP32 Sensors\n(field / reef)", bg="#059669")
    # Internet
    ax.text(3.3, 4.8, "☁ Internet", ha="center", va="center",
            fontsize=8, color="#6B7280")
    # Render
    render_rect = FancyBboxPatch((4.2, 3.2), 4.0, 3.0,
                                  boxstyle="round,pad=0.05",
                                  linewidth=2, edgecolor="#1D4ED8",
                                  facecolor="#EFF6FF", zorder=1)
    ax.add_patch(render_rect)
    ax.text(6.2, 6.0, "Render (Singapore)", fontsize=9, fontweight="bold",
            ha="center", color="#1D4ED8", zorder=3)
    box(ax, 6.2, 5.5, 3.3, 0.55, "Docker Container\n(python:3.11-slim-bookworm)", bg="#2563EB", fontsize=8)
    box(ax, 6.2, 4.8, 3.3, 0.55, "FastAPI + Uvicorn\n(port $PORT / 10000)", bg="#1D4ED8", fontsize=8)
    box(ax, 6.2, 4.1, 3.3, 0.55, "PINN + ANN-LSTM\n(bundled in model/)", bg="#7C3AED", fontsize=8)
    box(ax, 6.2, 3.5, 3.3, 0.4,  "APScheduler  (6h / 24h jobs)", bg="#6366F1", fontsize=7.5)

    # Supabase
    box(ax, 11.0, 5.0, 2.0, 0.65, "Supabase\n(PostgreSQL · SSL)", bg="#0F766E")
    # Vercel
    box(ax, 11.0, 3.8, 2.0, 0.65, "Vercel\n(Next.js Dashboard)", bg="#7C3AED")
    # GitHub
    box(ax, 6.2, 1.5, 3.3, 0.55, "GitHub\n(TeamTerronix/backend)", bg="#374151")
    box(ax, 6.2, 0.7, 3.3, 0.55, "GitHub Actions CI\n(pytest on push / PR)", bg="#374151")

    # Arrows
    arrow(ax, 2.1,  4.8, 4.2,  4.8, "HTTPS POST /data")
    arrow(ax, 8.2,  4.8, 9.0,  5.0, "SQL + SSL")
    arrow(ax, 9.0,  3.8, 8.2,  4.5, "reads API")
    arrow(ax, 6.2,  3.2, 6.2,  2.05, "Render deploys")
    arrow(ax, 6.2,  1.5, 6.2,  1.25)
    arrow(ax, 6.2,  0.45, 4.2, 3.5,  "deploy on push\nto main / host_in_render",
          color="#059669")

    # env note
    ax.add_patch(patches.FancyBboxPatch((0.1, 0.1), 3.5, 1.8,
                  boxstyle="round,pad=0.05", facecolor="#FEF3C7",
                  edgecolor="#D97706", lw=1))
    ax.text(1.85, 1.05,
            "Env vars (Render):\n  DATABASE_URL (Supabase)\n  SECRET_KEY\n  CORS_ORIGINS (Vercel URL)\n  DISABLE_SCHEDULER (opt.)",
            ha="center", va="center", fontsize=7, color="#92400E")

    return save(fig, "08_deployment.png")


# ═════════════════════════════════════════════════════════════════════════════
# DIAGRAM 9 — CI / Test Structure
# ═════════════════════════════════════════════════════════════════════════════

def diagram_ci():
    fig, ax = plt.subplots(figsize=(12, 5.5))
    ax.set_xlim(0, 12); ax.set_ylim(0, 5.5)
    ax.axis("off")
    fig.patch.set_facecolor("#F8FAFC")
    ax.set_facecolor("#F8FAFC")

    ax.text(6, 5.2, "SLIOT — Automated Testing & CI Pipeline",
            fontsize=12, fontweight="bold", ha="center", color="#1e3a5f")

    # Test files
    tests = [
        ("test_health.py",          "Health & public surface\n(GET /  · GET /ws/alerts HTTP probe)"),
        ("test_auth_unit.py",       "Auth unit tests\n(hash · verify · JWT roundtrip)"),
        ("test_auth_api.py",        "Auth API\n(register · login · /me · duplicate · 401)"),
        ("test_sensors_data.py",    "Sensors & ingest\n(list · register · POST /data · readings)"),
        ("test_lstm_api.py",        "ANN-LSTM endpoints\n(mocked TF · 503 · sensor lstm)"),
        ("test_database_url.py",    "DB URL helpers\n(normalize · mask · Supabase detect)"),
        ("test_scheduler_helpers.py","Scheduler helpers\n(location inference)"),
    ]

    colors = ["#1D4ED8","#0891B2","#6366F1","#059669","#D97706","#0F766E","#374151"]

    for i, ((fname, desc), col) in enumerate(zip(tests, colors)):
        x = 0.4 + (i % 4) * 2.9
        y = 3.8 if i < 4 else 1.8
        h = 1.0
        rect = FancyBboxPatch((x, y), 2.4, h,
                               boxstyle="round,pad=0.05",
                               facecolor=col, edgecolor="#1e3a5f",
                               lw=1.2, zorder=2)
        ax.add_patch(rect)
        ax.text(x+1.2, y+0.75, fname, ha="center", va="center",
                fontsize=7.5, color="white", fontweight="bold", zorder=3)
        ax.text(x+1.2, y+0.35, desc, ha="center", va="center",
                fontsize=6.5, color="#F3F4F6", zorder=3, multialignment="center")
        # arrow to CI
        ax.annotate("", xy=(6, 0.85), xytext=(x+1.2, y),
                    arrowprops=dict(arrowstyle="-|>", color="#9CA3AF",
                                    lw=0.8, connectionstyle="arc3,rad=0.0"), zorder=1)

    # CI box
    ci_rect = FancyBboxPatch((3.8, 0.3), 4.4, 0.7,
                              boxstyle="round,pad=0.05",
                              facecolor="#374151", edgecolor="#111827", lw=1.5, zorder=3)
    ax.add_patch(ci_rect)
    ax.text(6, 0.65,
            "GitHub Actions  ·  ubuntu-latest  ·  Python 3.11\npytest -q  (29 tests · SQLite in-memory · LSTM mocked)",
            ha="center", va="center", fontsize=7.5, color="white",
            fontweight="bold", zorder=4)

    return save(fig, "09_ci_tests.png")


# ═════════════════════════════════════════════════════════════════════════════
# BUILD WORD DOCUMENT
# ═════════════════════════════════════════════════════════════════════════════

def build_docx(img_paths: dict):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import datetime

    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Style helpers ──────────────────────────────────────────────────────────
    def h1(text):
        p = doc.add_heading(text, level=1)
        p.runs[0].font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        return p

    def h2(text):
        p = doc.add_heading(text, level=2)
        p.runs[0].font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        return p

    def h3(text):
        return doc.add_heading(text, level=3)

    def para(text, bold=False, italic=False, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(10.5)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    def bullet(text, level=0):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(level * 0.5)
        run = p.add_run(text)
        run.font.size = Pt(10)
        return p

    def table_row(tbl, cells, bold=False, bg=None):
        row = tbl.add_row()
        for i, txt in enumerate(cells):
            c = row.cells[i]
            c.text = txt
            for run in c.paragraphs[0].runs:
                run.bold = bold
                run.font.size = Pt(9)
            if bg:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), bg)
                shading.set(qn("w:val"), "clear")
                c._tc.get_or_add_tcPr().append(shading)
        return row

    def insert_image(path, width=Inches(5.8), caption=""):
        doc.add_picture(str(path), width=width)
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].font.size   = Pt(8.5)
            cp.runs[0].font.italic = True
            cp.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    def page_break():
        doc.add_page_break()

    def divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)

    # ══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("SLIOT — Benthic Guardian")
    r.font.size  = Pt(28)
    r.bold       = True
    r.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run("Backend Technical Documentation")
    r2.font.size  = Pt(16)
    r2.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    doc.add_paragraph()
    team_p = doc.add_paragraph()
    team_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = team_p.add_run("Team Terronix  ·  Smart Lankan IoT 2026")
    r3.font.size = Pt(12)
    r3.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = date_p.add_run(datetime.date.today().strftime("%B %Y"))
    r4.font.size  = Pt(11)
    r4.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()
    doc.add_paragraph()
    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rv = ver_p.add_run("Backend Version 1.0  ·  FastAPI · PostgreSQL · TensorFlow · Docker · Render")
    rv.font.size  = Pt(9)
    rv.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. PROJECT OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    h1("1. Project Overview")
    para(
        "SLIOT (Smart Lankan IoT) is a real-time coral reef monitoring platform built by Team Terronix. "
        "It collects water temperature data from ESP32 sensors placed at reef sites along Sri Lanka's "
        "coastline. The platform predicts future temperature changes and coral bleaching risk using two "
        "machine learning models, then shows the results in a live web dashboard."
    )
    doc.add_paragraph()
    para(
        "This document explains how the backend was designed, the decisions we made, and the technical "
        "details of each part. We use simple language so anyone on the team can understand it."
    )
    doc.add_paragraph()

    h2("1.1 What the System Does")
    bullet("Receives live temperature readings from underwater ESP32 sensors via HTTP POST.")
    bullet("Stores all readings and forecasts in a managed PostgreSQL database (Supabase).")
    bullet("Runs a Physics-Informed Neural Network (PINN) to forecast the next 7 days of temperature, every 6 hours.")
    bullet("Runs an ANN-LSTM model to forecast SST and bleaching risk at +1, +3, and +7 days ahead.")
    bullet("Sends real-time bleaching alerts to the dashboard via WebSocket when temperature exceeds 31°C.")
    bullet("Controls access using JWT tokens and role-based permissions (admin vs regular user).")

    doc.add_paragraph()
    h2("1.2 System Context")
    insert_image(img_paths["context"], Inches(5.8),
                 "Figure 1 — System Context: how actors and services connect to the backend")

    page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    h1("2. System Architecture")
    para(
        "The backend is a single FastAPI application running inside a Docker container on Render. "
        "It has several internal components that work together. The diagram below shows how they fit."
    )
    doc.add_paragraph()
    insert_image(img_paths["container"], Inches(6.0),
                 "Figure 2 — Container architecture: internal components and external dependencies")

    doc.add_paragraph()
    h2("2.1 Key Components")

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    table_row(tbl, ["Component", "File(s)", "What it does"], bold=True, bg="1D4ED8")
    rows = [
        ("FastAPI App",        "main.py",              "Registers all routes, lifespan startup, CORS, WebSocket manager"),
        ("Auth Module",        "auth.py",              "JWT creation/validation, bcrypt password hashing, role checks"),
        ("Database Layer",     "database.py",          "SQLAlchemy engine, session factory, URL normalization for Supabase"),
        ("ORM Models",         "models.py",            "Defines all 6 database tables as Python classes"),
        ("Scheduler",          "scheduler.py",         "APScheduler: runs PINN every 6h and data archival every 24h"),
        ("PINN Forecaster",    "model/forecaster.py",  "Loads pinn_model_best.h5, generates 168-hour hourly predictions"),
        ("ANN-LSTM",           "model/lstm_forecaster.py", "Loads ann_lstm_L60_best.h5, generates +1/+3/+7 day forecasts"),
        ("Model Paths",        "model_paths.py",       "Resolves model/ directory for both monorepo and standalone deploy"),
        ("Sync Script",        "sync_model_assets.py", "Copies model weights + datasets into backend/model/ for Docker"),
    ]
    for r in rows:
        table_row(tbl, list(r))

    doc.add_paragraph()
    h2("2.2 File Structure")
    para("The main files in the backend repository are:", bold=True)
    for line in [
        "main.py          —  All API routes, app factory, lifespan",
        "auth.py          —  JWT and bcrypt utilities",
        "database.py      —  Engine, session, URL handling",
        "models.py        —  SQLAlchemy ORM (6 tables)",
        "scheduler.py     —  Background jobs (PINN + archive)",
        "model_paths.py   —  Resolves model directory",
        "model/           —  ML weights, datasets, forecasters",
        "tests/           —  29 pytest tests",
        ".github/         —  GitHub Actions CI workflow",
        "Dockerfile        —  Docker image (python:3.11-slim-bookworm)",
        "render.yaml      —  Render blueprint configuration",
    ]:
        bullet(line)

    page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. DATABASE DESIGN
    # ══════════════════════════════════════════════════════════════════════════
    h1("3. Database Design")
    para(
        "The database runs on Supabase (managed PostgreSQL). There are 6 tables. "
        "SQLAlchemy is used as the ORM so the code does not write raw SQL."
    )
    doc.add_paragraph()
    insert_image(img_paths["erd"], Inches(6.5),
                 "Figure 3 — Entity Relationship Diagram (ERD): all tables and their relationships")

    doc.add_paragraph()
    h2("3.1 Table Descriptions")

    table_data = [
        ("users",                "Stores every account. Role is either 'admin' or 'user'. Password is stored as a bcrypt hash, never in plain text."),
        ("network_groups",       "A logical group of sensors (e.g. 'Hikkaduwa Reef Network'). ID is a string like 'ng_abc123'. Users belong to one or more groups."),
        ("user_network_groups",  "Join table connecting users and network groups. A user can be in many groups; a group can have many users (many-to-many)."),
        ("sensors",              "Each physical ESP32 device. Stores GPS coordinates, depth, the owner user, and a flag for approval. A sensor must be approved by an admin before it can POST data."),
        ("sensor_readings",      "One row per temperature reading received. No deduplication — full-resolution timestamps. Rows older than 30 days are archived to CSV and deleted."),
        ("predictions",          "PINN-generated 7-day (168 rows) forecast per sensor. Refreshed every 6 hours. Includes temperature, risk level (0/1/2), risk score (0–1), anomaly, and physics residual."),
    ]

    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.style = "Table Grid"
    table_row(tbl2, ["Table", "Description"], bold=True, bg="0891B2")
    for name, desc in table_data:
        table_row(tbl2, [name, desc])

    doc.add_paragraph()
    h2("3.2 Design Decisions")

    h3("Why no hourly deduplication?")
    para(
        "An earlier version had a unique constraint so only one reading per sensor per hour was stored. "
        "This caused 409 Conflict errors when sensors posted more frequently. We removed it "
        "(via migration) to store every reading at full resolution, which gives richer data for the PINN."
    )
    h3("Why Supabase session pooler (port 5432)?")
    para(
        "Supabase offers two pooler types. The transaction pooler (port 6543) does not support prepared "
        "statements, which SQLAlchemy uses by default. The session pooler (port 5432) behaves like a "
        "normal PostgreSQL connection and works without extra configuration. We also URL-encode special "
        "characters in the password to avoid connection string parsing errors."
    )
    h3("Why soft pool on Render?")
    para(
        "Render Starter has limited RAM. We set pool_size=5 and max_overflow=5 (10 total connections) "
        "to avoid overwhelming Supabase's free tier connection limit and to keep memory usage low."
    )

    page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. API DESIGN
    # ══════════════════════════════════════════════════════════════════════════
    h1("4. API Design")
    para(
        "The API is built with FastAPI. All endpoints return JSON. "
        "Most endpoints require a JWT bearer token. A few (health, register, login) are public."
    )
    doc.add_paragraph()

    h2("4.1 Endpoint Map")

    tbl3 = doc.add_table(rows=1, cols=4)
    tbl3.style = "Table Grid"
    table_row(tbl3, ["Method", "Path", "Auth", "Purpose"], bold=True, bg="1D4ED8")
    endpoints = [
        ("GET",  "/",                            "None",       "Health check — returns status, version, database state"),
        ("POST", "/auth/register",               "None",       "Create user account (role = 'user')"),
        ("POST", "/auth/token",                  "None",       "Login — returns JWT bearer token"),
        ("GET",  "/auth/me",                     "User",       "Return current user profile"),
        ("GET",  "/sensors",                     "User",       "List sensors (admin: all, user: own network)"),
        ("POST", "/admin/register-sensor",       "Admin",      "Register and approve a new sensor"),
        ("POST", "/admin/users",                 "Admin",      "Create a user account (admin-only)"),
        ("POST", "/admin/network-groups",        "Admin",      "Create a new network group"),
        ("GET",  "/admin/users",                 "Admin",      "List all user accounts"),
        ("POST", "/data",                        "None",       "Ingest sensor reading (sensor_uid + temperature)"),
        ("GET",  "/sensors/{id}/readings",       "User",       "Recent readings history for a sensor"),
        ("GET",  "/sensors/{id}/forecast",       "User",       "PINN forecast (168h) for a sensor"),
        ("GET",  "/sensors/{id}/lstm-forecast",  "User",       "ANN-LSTM forecast (+1/+3/+7d) for a sensor"),
        ("GET",  "/api/lstm-forecast",           "User",       "ANN-LSTM forecast by reef site location"),
        ("GET",  "/api/stats",                   "User",       "Dashboard KPIs (24h avg/max temp, counts)"),
        ("GET",  "/api/risk-summary",            "User",       "Risk counts (healthy/warning/danger) across network"),
        ("GET",  "/api/sst",                     "User",       "Sea Surface Temperature historical data"),
        ("GET",  "/api/dhw",                     "User",       "Degree Heating Week historical data"),
        ("GET",  "/api/latest-readings",         "User",       "Most recent reading per sensor"),
        ("WS",   "/ws/alerts",                   "JWT (query)","Push bleaching alerts when temp > 31°C"),
    ]
    for ep in endpoints:
        table_row(tbl3, list(ep))

    doc.add_paragraph()
    h2("4.2 Request Sequences")
    insert_image(img_paths["sequence"], Inches(6.2),
                 "Figure 4 — Key request sequences: login, sensor ingest, PINN forecast, LSTM forecast")

    page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. AUTHENTICATION
    # ══════════════════════════════════════════════════════════════════════════
    h1("5. Authentication & Authorization")
    para(
        "The backend uses JSON Web Tokens (JWT) for stateless authentication. "
        "Every protected endpoint uses a FastAPI dependency that decodes and validates the token."
    )
    doc.add_paragraph()
    insert_image(img_paths["auth"], Inches(5.8),
                 "Figure 5 — Authentication flow: login, bcrypt verify, JWT issue, role check")

    doc.add_paragraph()
    h2("5.1 How It Works")
    bullet("User sends email + password to POST /auth/token.")
    bullet("Backend looks up the user, then calls bcrypt.checkpw() to verify the password hash.")
    bullet("On success, it creates a JWT signed with SECRET_KEY (HS256). Token lifetime is 60 minutes.")
    bullet("Client stores the token and sends it as 'Authorization: Bearer <token>' on every request.")
    bullet("The get_current_user() dependency decodes the token and loads the user from the database.")
    bullet("The get_admin_user() dependency additionally checks that the role is 'admin'.")

    doc.add_paragraph()
    h2("5.2 Why bcrypt directly (not passlib)?")
    para(
        "We started with passlib for password hashing. However, passlib is incompatible with bcrypt "
        "version 4.x and above on Linux/Docker. When deployed to Render, every login and register "
        "returned a 500 Internal Server Error. We fixed this by calling bcrypt directly "
        "(bcrypt.hashpw and bcrypt.checkpw) which works correctly with any bcrypt version."
    )

    doc.add_paragraph()
    h2("5.3 Role-Based Access Control")
    tbl_r = doc.add_table(rows=1, cols=3)
    tbl_r.style = "Table Grid"
    table_row(tbl_r, ["Role", "Who", "What they can access"], bold=True, bg="6366F1")
    table_row(tbl_r, ["admin", "Team / project owner", "All endpoints, all sensors, all users, all networks"])
    table_row(tbl_r, ["user",  "Reef site owner",       "Only sensors and readings in their own network group"])

    page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 6. ML MODELS
    # ══════════════════════════════════════════════════════════════════════════
    h1("6. Machine Learning Models")
    para(
        "The backend integrates two machine learning models. They serve different purposes "
        "and complement each other. Both use TensorFlow and are loaded lazily to avoid "
        "slowing down the API start-up."
    )
    doc.add_paragraph()
    insert_image(img_paths["ml"], Inches(6.0),
                 "Figure 6 — ML models overview: PINN (spatial/hourly) vs ANN-LSTM (short-term sequence)")

    doc.add_paragraph()
    h2("6.1 PINN — Physics-Informed Neural Network")
    para(
        "The PINN was trained to predict ocean temperature at any (latitude, longitude, time) point "
        "around the reef sites. It uses the heat-equation PDE as a physics constraint in its training "
        "loss, which makes it physically realistic and better at extrapolating into the future."
    )
    bullet("Input: (latitude, longitude, time_hours)")
    bullet("Output: temperature prediction (°C) for that point and time")
    bullet("Forecast horizon: 168 hours (7 days) ahead, one row per hour")
    bullet("Trigger: every 6 hours by the scheduler, or immediately after a POST /data (throttled to once every 45 seconds)")
    bullet("Storage: predictions table in the database (old rows replaced on each run)")
    bullet("Files: model/pinn_model_best.h5, model/scalers.pkl, model/sensor_info.pkl")

    doc.add_paragraph()
    h2("6.2 ANN-LSTM — Short-Term Sequence Forecaster")
    para(
        "The ANN-LSTM uses a 60-day sliding window of Sea Surface Temperature (SST) and "
        "Degree Heating Weeks (DHW) to predict the next 1, 3, and 7 days. This model focuses "
        "on recent patterns rather than spatial physics, so it catches short-term warming trends "
        "that the PINN might smooth over."
    )
    bullet("Input: last 60 days of SST and DHW for a reef site")
    bullet("Output: predicted SST, DHW, anomaly, and bleaching risk at +1, +3, +7 days")
    bullet("Trigger: on API request (lazy-loaded — not stored in the DB)")
    bullet("Source data: bundled CSV history files (not live sensor_readings)")
    bullet("Locations: hikkaduwa, kalpitiya, passikudha, south_east, trinco")
    bullet("Files: model/ann_lstm_L60_best.h5, model/sliot_dataset/<location>/sst_full.csv")

    doc.add_paragraph()
    h2("6.3 Why Two Models?")
    tbl_m = doc.add_table(rows=1, cols=3)
    tbl_m.style = "Table Grid"
    table_row(tbl_m, ["", "PINN", "ANN-LSTM"], bold=True, bg="374151")
    table_row(tbl_m, ["What it predicts", "Temperature at any reef location", "SST + DHW for 5 named sites"])
    table_row(tbl_m, ["Time range",        "7 days (168 h hourly)",           "+1, +3, +7 days"])
    table_row(tbl_m, ["Data source",       "Live sensor readings",             "60-day CSV history"])
    table_row(tbl_m, ["Physics constraint","Yes (heat equation PDE)",          "No (data-driven only)"])
    table_row(tbl_m, ["Good at",           "Spatial maps, per-sensor forecast","Short-term SST trend changes"])
    table_row(tbl_m, ["Stored in DB?",     "Yes (predictions table)",          "No (computed on demand)"])

    page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 7. SCHEDULER & BACKGROUND JOBS
    # ══════════════════════════════════════════════════════════════════════════
    h1("7. Background Scheduler")
    para(
        "The backend runs two background jobs using APScheduler. They start when the API starts "
        "and run continuously while the server is up."
    )
    doc.add_paragraph()
    insert_image(img_paths["scheduler"], Inches(6.0),
                 "Figure 7 — Scheduler timeline: Job 1 (PINN every 6h), Job 2 (archive every 24h), and on-demand trigger")

    doc.add_paragraph()
    h2("7.1 Job 1 — PINN Forecast (every 6 hours)")
    bullet("Fetches the last 48 hours of readings for every approved sensor.")
    bullet("Runs the PINN model for each sensor to generate a 168-hour forecast.")
    bullet("Deletes old predictions for that sensor and writes fresh ones to the database.")
    bullet("If a sensor has no readings in the last 48 hours, it is skipped.")
    bullet("Also triggered by POST /data, but no more than once every 45 seconds.")

    doc.add_paragraph()
    h2("7.2 Job 2 — Data Archival (every 24 hours)")
    bullet("Finds sensor_readings rows older than 30 days.")
    bullet("Appends them to history_archive.csv (on disk).")
    bullet("Deletes them from the database to keep the 'hot' table fast.")
    bullet("If the CSV does not exist, it creates it with a header.")

    doc.add_paragraph()
    h2("7.3 Disabling the Scheduler")
    para(
        "Set the environment variable DISABLE_SCHEDULER=1 to skip APScheduler on startup. "
        "This is useful when running tests (done automatically via conftest.py) or when the "
        "Render instance is too small to handle TensorFlow in-process. "
        "The API continues to serve all non-forecast endpoints normally."
    )

    page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 8. WEBSOCKET ALERTS
    # ══════════════════════════════════════════════════════════════════════════
    h1("8. Real-Time Bleaching Alerts (WebSocket)")
    para(
        "When a sensor posts a temperature at or above 31°C (the coral bleaching threshold), "
        "the backend immediately broadcasts a bleaching_alert message to all connected WebSocket clients."
    )
    bullet("Clients connect to /ws/alerts?token=<jwt> using the wss:// protocol.")
    bullet("Admin users receive alerts from every sensor.")
    bullet("Regular users receive alerts only from sensors in their own network group.")
    bullet("If the client sends a plain HTTP GET to /ws/alerts (no Upgrade header), a 200 response with a helpful message is returned.")
    bullet("All connected clients also receive a reading_new event on every new POST /data.")

    page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 9. DEPLOYMENT
    # ══════════════════════════════════════════════════════════════════════════
    h1("9. Deployment")
    para(
        "The backend is deployed as a Docker container on Render (Singapore region). "
        "The database runs on Supabase (managed PostgreSQL). The dashboard is on Vercel."
    )
    doc.add_paragraph()
    insert_image(img_paths["deployment"], Inches(6.0),
                 "Figure 8 — Production deployment architecture: Render, Supabase, Vercel, GitHub")

    doc.add_paragraph()
    h2("9.1 How Deployment Works")
    bullet("Code is pushed to GitHub (TeamTerronix/backend-Benthic-Guardian).")
    bullet("Render picks up the push on the configured branch (main or host_in_render).")
    bullet("It builds the Docker image using Dockerfile (python:3.11-slim-bookworm base).")
    bullet("On startup, docker-entrypoint.sh checks the DB connection, creates tables if needed, then starts uvicorn on $PORT (default 10000).")
    bullet("Render scans for an open port. If it finds one within ~90 seconds, the deploy succeeds.")

    doc.add_paragraph()
    h2("9.2 Environment Variables")
    tbl_e = doc.add_table(rows=1, cols=3)
    tbl_e.style = "Table Grid"
    table_row(tbl_e, ["Variable", "Required?", "Description"], bold=True, bg="374151")
    envs = [
        ("DATABASE_URL",                "Yes",  "Supabase session pooler URI (port 5432). URL-encode password special chars."),
        ("SECRET_KEY",                  "Yes",  "Long random string for signing JWT tokens. Never share or commit."),
        ("ALGORITHM",                   "No",   "JWT algorithm. Defaults to HS256."),
        ("ACCESS_TOKEN_EXPIRE_MINUTES", "No",   "Token lifetime. Defaults to 60 minutes."),
        ("CORS_ORIGINS",                "Yes",  "Comma-separated list of allowed frontend URLs (e.g. your Vercel URL)."),
        ("CORS_ORIGIN_REGEX",           "No",   "Regex for CORS origins (e.g. https://.*\\.vercel\\.app)."),
        ("CUDA_VISIBLE_DEVICES",        "No",   "Set to -1 to force TensorFlow to use CPU (important on Render)."),
        ("DISABLE_SCHEDULER",           "No",   "Set to 1 to skip APScheduler. Useful for small instances or debugging."),
        ("PORT",                        "No",   "Port for uvicorn. Render sets this automatically."),
    ]
    for ev in envs:
        table_row(tbl_e, list(ev))

    doc.add_paragraph()
    h2("9.3 Known Deployment Issues")
    h3("Port scan timeout")
    para(
        "If the API process does not open a port within ~90 seconds, Render reports 'No open ports detected'. "
        "This usually means TensorFlow or the PINN is loading at startup and taking too long. "
        "Fix: set DISABLE_SCHEDULER=1 so the model does not load at boot."
    )
    h3("OOM crash (out of memory)")
    para(
        "TensorFlow + PINN and ANN-LSTM together use significant RAM. On Render Starter or Free tier, "
        "the process can be killed by the OS. This appears as a random crash with no Python traceback, "
        "just 'Killed' in the logs. Fix: upgrade instance size, or run forecasts in a separate worker service."
    )
    h3("Supabase paused")
    para(
        "The free Supabase project pauses after 7 days of inactivity. The API will fail all DB operations "
        "until you resume it from the Supabase dashboard. Redeploy or restart the Render service after resuming."
    )

    page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 10. TESTING & CI
    # ══════════════════════════════════════════════════════════════════════════
    h1("10. Testing & Continuous Integration")
    para(
        "The backend has 29 automated tests written with pytest. All tests run in CI on every "
        "push and pull request via GitHub Actions."
    )
    doc.add_paragraph()
    insert_image(img_paths["ci"], Inches(6.0),
                 "Figure 9 — Test structure and GitHub Actions CI pipeline")

    doc.add_paragraph()
    h2("10.1 Test Design Decisions")
    h3("SQLite in-memory for speed")
    para(
        "Tests use a shared in-memory SQLite database instead of the real Supabase. "
        "This makes tests fast (under 10 seconds) and does not require any network access. "
        "The DATABASE_URL is overridden to 'sqlite://' before the app imports."
    )
    h3("Mocked LSTM / PINN")
    para(
        "TensorFlow is not installed in the test environment (requirements-dev.txt has no TF). "
        "The _get_lstm() function in main.py is monkeypatched with a MagicMock that returns "
        "a fixed list of forecast points. This lets us test the API response shape without "
        "loading any model weights."
    )
    h3("DISABLE_SCHEDULER=1")
    para(
        "conftest.py sets this env var so APScheduler does not start, and no PINN jobs run "
        "during tests. The _schedule_forecast_job_after_reading() function is also patched "
        "to a no-op on the client fixture."
    )

    doc.add_paragraph()
    h2("10.2 Test Summary")
    tbl_t = doc.add_table(rows=1, cols=3)
    tbl_t.style = "Table Grid"
    table_row(tbl_t, ["Test file", "Count", "What is tested"], bold=True, bg="059669")
    test_rows = [
        ("test_health.py",           "2",  "GET / returns ok + database ok. /ws/alerts HTTP probe returns 200."),
        ("test_auth_unit.py",        "3",  "hash_password, verify_password, JWT encode/decode roundtrip."),
        ("test_auth_api.py",         "5",  "Register, login, /me, duplicate email 400, wrong password 401."),
        ("test_sensors_data.py",     "8",  "List sensors, admin register, unapproved block, ingest, readings history."),
        ("test_lstm_api.py",         "4",  "LSTM requires auth, mock forecast, 503 on missing weights, per-sensor."),
        ("test_database_url.py",     "4",  "URL normalize, quote-strip, Supabase detect, mask (password hidden)."),
        ("test_scheduler_helpers.py","3",  "Location inference from UID prefix and from GPS coordinates."),
    ]
    for tr in test_rows:
        table_row(tbl_t, list(tr))

    doc.add_paragraph()
    h2("10.3 Running Tests Locally")
    para("Install dev dependencies and run pytest:", bold=True)
    code_p = doc.add_paragraph()
    code_r = code_p.add_run(
        "pip install -r requirements-dev.txt\n"
        "pytest -q"
    )
    code_r.font.name = "Courier New"
    code_r.font.size = Pt(9)

    page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 11. KEY DESIGN DECISIONS
    # ══════════════════════════════════════════════════════════════════════════
    h1("11. Key Design Decisions")

    decisions = [
        ("FastAPI over Flask or Django",
         "FastAPI gives automatic OpenAPI docs, Pydantic validation, async support, and is faster. "
         "Django was too heavy for a pure API service. Flask lacks built-in validation."),
        ("bcrypt directly (not passlib)",
         "passlib's bcrypt wrapper is incompatible with bcrypt>=4.0.0 on Linux/Docker (used on Render). "
         "Calling bcrypt.hashpw() and bcrypt.checkpw() directly works with any version."),
        ("JWT (stateless) over session cookies",
         "The dashboard and ESP32 sensors both need auth. Stateless JWT works across domains without "
         "a shared session store. Tokens expire after 60 minutes."),
        ("Supabase (managed PostgreSQL) over self-hosted",
         "Free managed database with automatic backups, SSL, and a pooler. Avoids running and paying "
         "for a separate database server. Matches Render's free/starter tier budget."),
        ("Session pooler port 5432 over transaction pooler 6543",
         "The transaction pooler disables prepared statements, which SQLAlchemy uses by default. "
         "Session pooler behaves like a normal PostgreSQL connection and needs no extra config."),
        ("Docker on Render over buildpack",
         "Docker gives full control over the runtime environment, including TensorFlow dependencies "
         "and model files. Buildpacks cannot reliably install TF or bundle large binary files."),
        ("PINN + ANN-LSTM (two models)",
         "PINN handles spatial mapping and per-sensor hourly forecasts using physics constraints. "
         "ANN-LSTM handles short-term SST/DHW trends using sequence data. They cover different "
         "time scales and data types so both are useful."),
        ("Lazy model loading",
         "Both models are loaded on first use, not at import time. This keeps API startup fast "
         "and avoids loading TF when it is not needed (e.g. during tests)."),
        ("Removing hourly deduplication",
         "The original unique constraint on (sensor_id, hour) caused 409 errors when sensors "
         "posted more than once per hour. Removing it stores full-resolution data, which is more "
         "accurate and flexible for ML training."),
        ("In-memory SQLite for tests",
         "Tests should not require a real database connection. SQLite in-memory is fast, isolated, "
         "and works on any machine including GitHub Actions CI runners with no DB server."),
    ]

    for i, (title, explanation) in enumerate(decisions, 1):
        h3(f"{i}. {title}")
        para(explanation)
        doc.add_paragraph()

    page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 12. KNOWN LIMITATIONS & FUTURE WORK
    # ══════════════════════════════════════════════════════════════════════════
    h1("12. Known Limitations & Future Work")

    h2("12.1 Current Limitations")
    bullet("TensorFlow uses a lot of RAM. Running both PINN and ANN-LSTM on a Render Starter instance can cause OOM crashes.")
    bullet("The scheduler runs inside the web process. Heavy forecast jobs can delay API responses.")
    bullet("ANN-LSTM uses bundled CSV history, not live sensor_readings. It will not reflect brand-new sensor sites without retraining.")
    bullet("WebSocket connections are stored in memory. If the server restarts, all connected clients lose their connection.")
    bullet("No rate limiting on POST /data. A malfunctioning sensor could flood the database.")

    doc.add_paragraph()
    h2("12.2 Recommended Future Improvements")
    bullet("Move forecast jobs to a separate Render background worker to isolate TF memory from the web API.")
    bullet("Add rate limiting (e.g. slowapi) on sensor data ingestion.")
    bullet("Persist WebSocket sessions via Redis for multi-instance deployments.")
    bullet("Retrain ANN-LSTM periodically on new sensor data to keep forecasts accurate.")
    bullet("Add pagination to /sensors/{id}/readings for large datasets.")
    bullet("Add structured logging with request IDs for easier debugging on Render.")

    page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 13. RELATED REPOS & LINKS
    # ══════════════════════════════════════════════════════════════════════════
    h1("13. Related Repositories & Links")
    tbl_l = doc.add_table(rows=1, cols=2)
    tbl_l.style = "Table Grid"
    table_row(tbl_l, ["Item", "Link / Notes"], bold=True, bg="374151")
    links = [
        ("Backend repo",          "github.com/TeamTerronix/backend-Benthic-Guardian"),
        ("Dashboard repo",        "Next.js frontend (separate repo)"),
        ("Model repo",            "PINN + ANN-LSTM training notebooks (separate repo)"),
        ("Live API",              "https://backend-benthic-guardian.onrender.com"),
        ("API docs (Swagger)",    "https://backend-benthic-guardian.onrender.com/docs"),
        ("Database",              "Supabase (Singapore) — session pooler port 5432"),
        ("Deployment guide",      "RENDER_DEPLOYMENT.md in the backend repo"),
    ]
    for lnk in links:
        table_row(tbl_l, list(lnk))

    # ══════════════════════════════════════════════════════════════════════════
    # SAVE
    # ══════════════════════════════════════════════════════════════════════════
    out_path = DOCS_DIR / "SLIOT_Backend_Documentation.docx"
    doc.save(str(out_path))
    print(f"\n  Word document saved -> {out_path}")
    return out_path


# ═════════════════════════════════════════════════════════════════════════════
# MAIN
# ═════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("Generating diagrams...")
    img_paths = {
        "context":    diagram_system_context(),
        "container":  diagram_container(),
        "erd":        diagram_erd(),
        "sequence":   diagram_sequence(),
        "scheduler":  diagram_scheduler(),
        "auth":       diagram_auth(),
        "ml":         diagram_ml(),
        "deployment": diagram_deployment(),
        "ci":         diagram_ci(),
    }

    print("\nBuilding Word document...")
    out = build_docx(img_paths)

    print(f"\nDone! Open: {out}")
